import os
import tempfile
from pathlib import Path
from typing import Tuple

import win32api
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .word_feature import WordFeature
from ..entity.base.experience_base import ProjectExperience, WorkExperience
from ..entity.base.info_base import InfoBase
from ..entity.info.basic_info import BasicInfo
from ..util.import_util import ImportUtil
from ..util.word_util import WordUtil


class ResumeFeature:
    ImportUtil.import_module(Path(__file__).parent.parent.joinpath("entity"))  # 使用subclasses之前必须将子类导入
    _resume_path = None

    @classmethod
    def add_title(cls):
        """添加标题"""
        document = WordFeature.get_document()
        paragraph = document.add_paragraph()
        WordUtil.set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
        run = paragraph.add_run("个人简历")
        WordUtil.set_bold(run)
        WordUtil.set_font_type(run, "黑体")
        WordUtil.set_font_size(run, 26)  # 一号字体

    @classmethod
    def add_info(cls):
        """添加信息"""
        # 1) 添加标题
        cls._add_heading_1("基本信息")
        # 2) 添加段落
        for info_class in InfoBase.__subclasses__():
            info = info_class()
            # 2.1) 添加信息名称
            if not isinstance(info, BasicInfo):
                cls._add_heading_2(info.info_name)
            for context in info.to_contexts():
                cls._add_paragraph(20, context)

    @classmethod
    def add_work_experience(cls):
        """添加工作经验"""
        # 1) 添加标题
        cls._add_heading_1("工作经验")
        # 2) 添加段落
        for work_experience_class in WorkExperience.__subclasses__()[::-1]:
            work_experience = work_experience_class()
            # 2.1) 添加公司名称
            company_name = work_experience.get_company_name()
            date_range = work_experience.get_date_range()
            cls._add_heading_2(company_name, date_range)
            # 2.2) 添加工作描述
            cls._add_paragraph(20, work_experience.get_profile())
            cls._add_paragraph(20, f"职位: {work_experience.get_job_position()}")
            cls._add_paragraph(20, "工作内容: ")
            for work_detail in work_experience.get_work_details():
                cls._add_paragraph(30, work_detail)

    @classmethod
    def add_project_experience(cls):
        """添加项目经验"""
        # 1) 添加标题
        cls._add_heading_1("项目经验")
        # 2) 添加段落
        for project_experience_class in ProjectExperience.__subclasses__()[::-1]:
            project_experience = project_experience_class()
            # 2.1) 添加项目名称
            project_name = project_experience.get_project_name()
            date_range = project_experience.get_date_range()
            cls._add_heading_2(project_name, date_range)
            # 2.2) 添加项目描述
            cls._add_paragraph(20, project_experience.get_profile())
            cls._add_paragraph(20, "开发技术: %s" % "、".join(project_experience.get_technologies()))
            cls._add_paragraph(20, "项目内容: ")
            for project_detail in project_experience.get_project_details():
                cls._add_paragraph(30, project_detail)

    @classmethod
    def resume_show(cls):
        resume_path = cls._get_resume_path()
        WordFeature.save_document(resume_path)
        win32api.ShellExecute(0, "open", resume_path, "", "", 1)

    @staticmethod
    def _add_heading_1(heading_name: str):
        """添加标题"""
        document = WordFeature.get_document()
        heading = document.add_heading(level=1)
        WordUtil.set_paragraph_space_before(heading, 6)
        run = heading.add_run(heading_name)
        WordUtil.set_bold(run)
        WordUtil.set_font_type(run, "宋体")
        WordUtil.set_font_size(run, 22)  # 二号字体

    @staticmethod
    def _add_heading_2(heading_name: str, date_range: Tuple[str, str] = None):
        """添加标题日期"""
        document = WordFeature.get_document()
        heading = document.add_heading(level=2)
        WordUtil.set_paragraph_indent(heading, 10)
        WordUtil.set_paragraph_space_before(heading, 6)
        run = heading.add_run(heading_name)
        WordUtil.set_font_type(run, "宋体")
        WordUtil.set_font_size(run, 14)  # 四号字体
        if date_range:
            run = heading.add_run(" | %s-%s" % date_range)
            WordUtil.set_font_type(run, "宋体")
            WordUtil.set_font_size(run, 12)  # 小四号字体

    @staticmethod
    def _add_paragraph(indent: int, context: str):
        """添加段落"""
        document = WordFeature.get_document()
        paragraph = document.add_paragraph()
        WordUtil.set_paragraph_indent(paragraph, indent)
        WordUtil.set_paragraph_space_after(paragraph, 0)
        run = paragraph.add_run(context)
        WordUtil.set_font_type(run, "仿宋")
        WordUtil.set_font_size(run, 12)  # 小四号字体

    @classmethod
    def _get_resume_path(cls) -> str:
        if cls._resume_path is None:
            cls._resume_path = os.path.join(tempfile.mkdtemp(), "个人简历.docx")
        return cls._resume_path
