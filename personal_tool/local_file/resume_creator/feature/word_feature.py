from typing import Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from personal_tool.local_file.resume_creator.config.resume_word_config import ResumeWordConfig
from personal_tool.local_file.resume_creator.entity.base.experience_base import WorkExperience, ProjectExperience
from personal_tool.local_file.resume_creator.entity.base.info_base import InfoBase
from personal_tool.local_file.resume_creator.util.word_util import WordUtil


class WordFeature:

    @staticmethod
    def add_title(document: Document):
        """添加标题"""
        paragraph = document.add_paragraph()
        WordUtil.set_paragraph_alignment(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
        run = paragraph.add_run("个人简历")
        WordUtil.set_bold(run)
        WordUtil.set_font_type(run, ResumeWordConfig.title_font_type)
        WordUtil.set_font_size(run, ResumeWordConfig.title_font_size)

    @classmethod
    def add_info(cls, document: Document):
        """添加信息"""
        # 1) 添加标题
        cls._add_heading(document, "基本信息")
        # 2) 添加段落
        for info_class in InfoBase.__subclasses__():
            info = info_class()
            cls._add_paragraph(document, 10, info.to_text())

    @classmethod
    def add_work_experience(cls, document: Document):
        """添加工作经验"""
        # 1) 添加标题
        cls._add_heading(document, "工作经验")
        # 2) 添加段落
        for work_experience_class in WorkExperience.__subclasses__():
            work_experience = work_experience_class()
            # 2.1) 添加公司名称
            company_name = work_experience.get_company_name()
            date_range = work_experience.get_date_range()
            cls._add_heading_and_date(document, company_name, date_range)
            # 2.2) 添加工作描述
            cls._add_paragraph(document, 20, f"职位：{work_experience.get_post()}")
            cls._add_paragraph(document, 20, work_experience.get_profile())

    @classmethod
    def add_project_experience(cls, document: Document):
        """添加项目经验"""
        # 1) 添加标题
        cls._add_heading(document, "项目经验")
        # 2) 添加段落
        for project_experience_class in ProjectExperience.__subclasses__():
            project_experience = project_experience_class()
            # 2.1) 添加项目名称
            project_name = project_experience.get_project_name()
            date_range = project_experience.get_date_range()
            cls._add_heading_and_date(document, project_name, date_range)
            # 2.2) 添加项目描述
            cls._add_paragraph(document, 20, project_experience.get_profile())
            cls._add_paragraph(document, 20, "开发技术：%s" % "、".join(project_experience.get_technologies()))
            cls._add_paragraph(document, 20, f"项目内容：")
            for project_detail in project_experience.get_project_details():
                cls._add_paragraph(document, 30, project_detail)

    @staticmethod
    def _add_heading(document: Document, heading_name: str):
        """添加标题"""
        heading = document.add_heading(level=1)
        WordUtil.set_paragraph_space_before(heading, 6)
        run = heading.add_run(heading_name)
        WordUtil.set_bold(run)
        WordUtil.set_font_type(run, ResumeWordConfig.heading_font_type)
        WordUtil.set_font_size(run, ResumeWordConfig.heading_font_size_1)

    @staticmethod
    def _add_heading_and_date(document: Document, heading_name: str, date_range: Tuple[str, str]):
        """添加标题日期"""
        heading = document.add_heading(level=2)
        WordUtil.set_paragraph_indent(heading, 10)
        WordUtil.set_paragraph_space_before(heading, 6)
        run = heading.add_run(heading_name)
        WordUtil.set_font_type(run, ResumeWordConfig.heading_font_type)
        WordUtil.set_font_size(run, ResumeWordConfig.heading_font_size_2)
        run = heading.add_run("|%s-%s" % date_range)
        WordUtil.set_font_type(run, ResumeWordConfig.heading_font_type)
        WordUtil.set_font_size(run, ResumeWordConfig.body_font_size)

    @staticmethod
    def _add_paragraph(document: Document, indent: int, context: str):
        """添加段落"""
        paragraph = document.add_paragraph()
        WordUtil.set_paragraph_indent(paragraph, indent)
        WordUtil.set_paragraph_space_after(paragraph, 0)
        run = paragraph.add_run(context)
        WordUtil.set_font_type(run, ResumeWordConfig.body_font_type)
        WordUtil.set_font_size(run, ResumeWordConfig.body_font_size)
